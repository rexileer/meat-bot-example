import re
from datetime import datetime
from decimal import Decimal
from io import BytesIO

from docx import Document

from Web.CRM.models import RawMeatBatch, RawMeatBatchStatus
from data.config import env
from bot.utils.docx import docx_replace_regex
from bot.utils.helpers import beautify_decimal


def generate_acceptance_certificate(raw_meat_batch: RawMeatBatch) -> BytesIO:
    document = Document(env.APP_DIR / "files" / "acceptance_certificate.docx")

    date_format = "%d.%m.%Y"

    now = datetime.utcnow()
    now_date = now.strftime(date_format)
    docx_replace_regex(document, re.compile(r"now_date"), now_date)
    docx_replace_regex(
        document, re.compile(r"yearnumber"), str(RawMeatBatch.generate_raw_meat_batch_year_number(raw_meat_batch))
    )
    docx_replace_regex(document, re.compile(r"production_id"), raw_meat_batch.production_id)
    docx_replace_regex(
        document,
        re.compile(r"organization"),
        raw_meat_batch.organization_vet if raw_meat_batch.organization_vet else raw_meat_batch.organization,
    )
    docx_replace_regex(
        document, re.compile(r"manufacturer"), raw_meat_batch.manufacturer if raw_meat_batch.manufacturer else ""
    )
    docx_replace_regex(document, re.compile(r"rawmaterialname"), raw_meat_batch.raw_material.name)
    docx_replace_regex(
        document,
        re.compile(r"manufacture_date_vet"),
        raw_meat_batch.manufacture_date_vet if raw_meat_batch.manufacture_date_vet else "",
    )
    docx_replace_regex(
        document,
        re.compile(r"expiration_date_vet"),
        raw_meat_batch.expiration_date_vet if raw_meat_batch.expiration_date_vet else "",
    )
    docx_replace_regex(
        document,
        re.compile(r"documentnumbervet"),
        raw_meat_batch.document_number_vet if raw_meat_batch.document_number_vet else "",
    )
    docx_replace_regex(
        document,
        re.compile(r"document_date_vet"),
        raw_meat_batch.document_date_vet if raw_meat_batch.document_date_vet else "",
    )
    docx_replace_regex(
        document, re.compile(r"numberttn"), raw_meat_batch.number_ttn if raw_meat_batch.number_ttn else ""
    )
    docx_replace_regex(
        document,
        re.compile(r"date_ttn"),
        raw_meat_batch.date_ttn.strftime(date_format) if raw_meat_batch.date_ttn else "",
    )
    docx_replace_regex(
        document,
        re.compile(r"productiontemperature"),
        raw_meat_batch.temperature if raw_meat_batch.temperature else "",
    ),
    docx_replace_regex(document, re.compile(r"raw_meat_batch_weight"), beautify_decimal(raw_meat_batch.weight))

    analyze_status = RawMeatBatchStatus.objects.filter(
        status__codename="laboratory_analyze", raw_meat_batch=raw_meat_batch
    ).first()
    if analyze_status and analyze_status.additional_data:
        docx_replace_regex(
            document,
            re.compile(r"appearance"),
            "Соотв." if analyze_status.additional_data["appearance"] else "Не соотв.",
        )
        docx_replace_regex(
            document, re.compile(r"smell"), "Соотв." if analyze_status.additional_data["smell"] else "Не соотв."
        )
        docx_replace_regex(
            document, re.compile(r"color"), "Соотв." if analyze_status.additional_data["color"] else "Не соотв."
        )
        docx_replace_regex(
            document,
            re.compile(r"broth_quality"),
            (
                analyze_status.additional_data.get("broth_quality")
                if analyze_status.additional_data.get("broth_quality")
                else ""
            ),
        )
        docx_replace_regex(
            document,
            re.compile(r"broth_flavor"),
            "Соотв." if analyze_status.additional_data["broth_flavor"] else "Не соотв.",
        )
        docx_replace_regex(document, re.compile(r"betta_lactams"), analyze_status.additional_data["betta_lactams"])
        docx_replace_regex(
            document, re.compile(r"chloramphenicols"), analyze_status.additional_data["chloramphenicols"]
        )
        docx_replace_regex(document, re.compile(r"tetracyclines"), analyze_status.additional_data["tetracyclines"])
        docx_replace_regex(document, re.compile(r"streptomycins"), analyze_status.additional_data["streptomycins"])
        docx_replace_regex(
            document,
            re.compile(r"fat_proportion"),
            beautify_decimal(Decimal(analyze_status.additional_data["fat_proportion"])),
        )
        docx_replace_regex(
            document,
            re.compile(r"protein_proportion"),
            beautify_decimal(Decimal(analyze_status.additional_data["protein_proportion"])),
        )
        docx_replace_regex(
            document,
            re.compile(r"moisture_proportion"),
            beautify_decimal(Decimal(analyze_status.additional_data["moisture_proportion"])),
        )
        docx_replace_regex(document, re.compile(r"responsible"), analyze_status.additional_data["responsible"])
    else:
        docx_replace_regex(document, re.compile(r"appearance"), "")
        docx_replace_regex(document, re.compile(r"smell"), "")
        docx_replace_regex(document, re.compile(r"color"), "")
        docx_replace_regex(document, re.compile(r"broth_quality"), "")
        docx_replace_regex(document, re.compile(r"broth_flavor"), "")
        docx_replace_regex(document, re.compile(r"betta_lactams"), "")
        docx_replace_regex(document, re.compile(r"chloramphenicols"), "")
        docx_replace_regex(document, re.compile(r"tetracyclines"), "")
        docx_replace_regex(document, re.compile(r"streptomycins"), "")
        docx_replace_regex(document, re.compile(r"fat_proportion"), "")
        docx_replace_regex(document, re.compile(r"protein_proportion"), "")
        docx_replace_regex(document, re.compile(r"moisture_proportion"), "")
        docx_replace_regex(document, re.compile(r"responsible"), "")

    if analyze_status and analyze_status.additional_data.get("separator_name"):
        docx_replace_regex(
            document,
            re.compile(r"separator"),
            f"{analyze_status.additional_data['separator_name']} - {analyze_status.additional_data['separator_mode']}",
        )
    else:
        docx_replace_regex(document, re.compile(r"separator"), "")

    """
    tables = document.tables
    p = tables[1].rows[1].cells[1].add_paragraph()
    r = p.add_run()
    master = f"{DATA_DIR}/files/Аминджонов.png"
    r.add_picture(master, width=Inches(.50), height=Inches(.35))

    p = tables[1].rows[2].cells[1].add_paragraph()
    r = p.add_run()
    laboratory_employee = f"{DATA_DIR}/files/Макунина.png"
    r.add_picture(laboratory_employee, width=Inches(.50), height=Inches(.35))

    p = tables[1].rows[3].cells[1].add_paragraph()
    r = p.add_run()
    boss = f"{DATA_DIR}/files/Пронских.png"
    r.add_picture(boss, width=Inches(.50), height=Inches(.35))
    """

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream
